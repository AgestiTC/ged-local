"""
Service Export — Génération de fichiers PDF et DOCX
====================================================
Convertit le texte Markdown d'un rapport en PDF (weasyprint)
ou en DOCX (python-docx).

Utilisé par les routers d'export et potentiellement en batch.
"""

from datetime import datetime
from pathlib import Path

from logger import get_logger

log = get_logger(__name__)

# Template HTML minimal pour l'export PDF
HTML_TEMPLATE = """\
<!DOCTYPE html>
<html lang="fr">
<head>
  <meta charset="UTF-8">
  <title>{title}</title>
  <style>
    body {{
      font-family: 'Helvetica Neue', Arial, sans-serif;
      font-size: 11pt;
      line-height: 1.7;
      color: #1a1a1a;
      max-width: 780px;
      margin: 40px auto;
      padding: 0 24px;
    }}
    h1 {{ font-size: 22pt; color: #111827; border-bottom: 2px solid #e5e7eb; padding-bottom: 10px; margin-bottom: 24px; }}
    h2 {{ font-size: 16pt; color: #1f2937; margin-top: 32px; margin-bottom: 12px; }}
    h3 {{ font-size: 13pt; color: #374151; margin-top: 24px; margin-bottom: 8px; }}
    h4 {{ font-size: 11pt; color: #4b5563; font-weight: 600; }}
    p  {{ margin: 0 0 12px 0; }}
    ul, ol {{ margin: 0 0 12px 0; padding-left: 24px; }}
    li {{ margin-bottom: 4px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 16px 0; font-size: 10pt; }}
    th, td {{ border: 1px solid #d1d5db; padding: 8px 12px; text-align: left; }}
    th {{ background-color: #f3f4f6; font-weight: 600; }}
    tr:nth-child(even) td {{ background-color: #f9fafb; }}
    code {{ background: #f3f4f6; padding: 2px 5px; border-radius: 3px; font-size: 10pt; font-family: 'Courier New', monospace; }}
    pre  {{ background: #f3f4f6; padding: 16px; border-radius: 6px; overflow-x: auto; font-size: 9pt; }}
    blockquote {{ border-left: 4px solid #3b82f6; padding-left: 16px; color: #4b5563; margin: 16px 0; font-style: italic; }}
    hr {{ border: none; border-top: 1px solid #e5e7eb; margin: 24px 0; }}
    .footer {{ margin-top: 48px; padding-top: 12px; border-top: 1px solid #e5e7eb; font-size: 9pt; color: #9ca3af; text-align: right; }}
    @page {{ margin: 2.5cm 2cm; }}
    @page :first {{ margin-top: 3cm; }}
  </style>
</head>
<body>
  <h1>{title}</h1>
  {content}
  <div class="footer">
    Généré par DocFlow AI — {date}
  </div>
</body>
</html>"""


def _nom_fichier(title: str, extension: str) -> str:
    """Génère un nom de fichier propre (sans caractères spéciaux)."""
    safe = "".join(c if c.isalnum() or c in "-_ " else "_" for c in title)
    safe = safe.strip().replace(" ", "_")[:50]
    horodatage = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{safe}_{horodatage}.{extension}"


class ExportService:
    """Gère l'export des rapports en PDF et DOCX."""

    def __init__(self, exports_dir: str):
        self.exports_dir = Path(exports_dir)
        self.exports_dir.mkdir(parents=True, exist_ok=True)

    async def to_pdf(self, content: str, title: str) -> Path:
        """
        Convertit du Markdown en PDF via weasyprint.

        Args:
            content: Contenu Markdown du rapport
            title: Titre du rapport

        Returns:
            Chemin vers le fichier PDF généré
        """
        try:
            import markdown
            from weasyprint import HTML
        except ImportError as e:
            raise RuntimeError(f"Dépendance manquante pour export PDF : {e}") from e

        nom_fichier = _nom_fichier(title, "pdf")
        chemin_pdf = self.exports_dir / nom_fichier

        # Markdown → HTML
        contenu_html = markdown.markdown(
            content,
            extensions=["tables", "fenced_code", "nl2br", "sane_lists"],
        )

        html_complet = HTML_TEMPLATE.format(
            title=title,
            content=contenu_html,
            date=datetime.now().strftime("%d/%m/%Y à %H:%M"),
        )

        try:
            HTML(string=html_complet).write_pdf(str(chemin_pdf))
        except Exception as e:
            log.error("Erreur génération PDF weasyprint", erreur=str(e))
            raise RuntimeError(f"Erreur génération PDF : {e}") from e

        log.info("PDF exporté", fichier=nom_fichier, taille=chemin_pdf.stat().st_size)
        return chemin_pdf

    async def to_docx(self, content: str, title: str) -> Path:
        """
        Convertit du Markdown en DOCX via python-docx.

        Args:
            content: Contenu Markdown du rapport
            title: Titre du rapport

        Returns:
            Chemin vers le fichier DOCX généré
        """
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, RGBColor
        except ImportError as e:
            raise RuntimeError(f"Dépendance manquante pour export DOCX : {e}") from e

        nom_fichier = _nom_fichier(title, "docx")
        chemin_docx = self.exports_dir / nom_fichier

        doc = DocxDocument()

        # Titre principal
        titre = doc.add_heading(title, level=0)
        if titre.runs:
            titre.runs[0].font.size = Pt(22)

        # Traitement ligne par ligne du Markdown
        lignes = content.split("\n")
        i = 0

        while i < len(lignes):
            ligne = lignes[i]

            if ligne.startswith("#### "):
                doc.add_heading(ligne[5:].strip(), level=4)
            elif ligne.startswith("### "):
                doc.add_heading(ligne[4:].strip(), level=3)
            elif ligne.startswith("## "):
                doc.add_heading(ligne[3:].strip(), level=2)
            elif ligne.startswith("# "):
                doc.add_heading(ligne[2:].strip(), level=1)
            elif ligne.startswith(("- ", "* ", "+ ")):
                texte = ligne[2:].strip()
                if texte:
                    doc.add_paragraph(texte, style="List Bullet")
            elif ligne.startswith("> "):
                p = doc.add_paragraph(ligne[2:].strip())
                p.paragraph_format.left_indent = Pt(24)
                if p.runs:
                    p.runs[0].font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
                    p.runs[0].font.italic = True
            elif ligne.startswith("---") and len(ligne.strip()) >= 3 and all(c in "-" for c in ligne.strip()):
                # Séparateur horizontal — ajouter un paragraphe vide avec bordure
                doc.add_paragraph()
            elif ligne.strip() == "":
                pass
            else:
                texte = ligne.strip()
                if texte:
                    doc.add_paragraph(texte)

            i += 1

        # Pied de page
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        run = p.add_run(f"Généré par DocFlow AI — {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        try:
            doc.save(str(chemin_docx))
        except Exception as e:
            log.error("Erreur génération DOCX", erreur=str(e))
            raise RuntimeError(f"Erreur génération DOCX : {e}") from e

        log.info("DOCX exporté", fichier=nom_fichier, taille=chemin_docx.stat().st_size)
        return chemin_docx
