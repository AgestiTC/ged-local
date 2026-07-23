"""
Router Export — /api/export
============================
Export de contenu Markdown en PDF ou DOCX.

Endpoints :
  POST /export/pdf    → Markdown → PDF (weasyprint)
  POST /export/docx   → Markdown → DOCX (python-docx)
"""

from datetime import datetime
from io import BytesIO
from urllib.parse import quote

import markdown
from docx import Document as DocxDocument
from docx.shared import Pt
from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel, Field

from logger import get_logger

log = get_logger(__name__)
router = APIRouter()


def _fichier_reponse(data: bytes, nom_fichier: str, media_type: str) -> Response:
    """
    Renvoie des octets en pièce jointe téléchargeable, SANS écrire sur disque.
    Évite toute dépendance aux droits du montage `storage/exports` (le conteneur tourne en
    uid 10001 ; un montage root donnait « Permission denied » au `doc.save()`). `filename*`
    encode l'UTF-8 pour les accents.
    """
    dispo = f"attachment; filename*=UTF-8''{quote(nom_fichier)}"
    return Response(content=data, media_type=media_type, headers={"Content-Disposition": dispo})


# Feuille de style du PDF — rendu « document » soigné (typographie, titres à accent,
# tableaux zébrés, encart Sources, pied de page paginé). Auto-suffisante (WeasyPrint : aucune
# ressource externe). Palette sobre indigo/ardoise.
_PDF_CSS = """
  @page {
    size: A4; margin: 2cm 1.8cm 2.2cm;
    @bottom-center {
      content: "__TITRE_COURT__";
      font-family: 'DejaVu Sans', sans-serif; font-size: 7.5pt; color: #9ca3af;
    }
    @bottom-right {
      content: "Page " counter(page) " / " counter(pages);
      font-family: 'DejaVu Sans', sans-serif; font-size: 7.5pt; color: #9ca3af;
    }
  }
  * { box-sizing: border-box; }
  body {
    font-family: 'DejaVu Sans', 'Helvetica Neue', Arial, sans-serif;
    font-size: 10.5pt; line-height: 1.62; color: #1f2937; margin: 0;
  }
  /* En-tête du document */
  .doc-header { border-left: 5px solid #4f46e5; padding: 2px 0 2px 14px; margin-bottom: 22px; }
  .doc-header h1 { font-size: 21pt; color: #1e1b4b; margin: 0 0 4px; line-height: 1.2; }
  .doc-meta { font-size: 8.5pt; color: #6b7280; text-transform: uppercase; letter-spacing: .06em; }
  /* Titres du contenu */
  h1 { font-size: 16pt; color: #312e81; margin: 22px 0 8px; }
  h2 {
    font-size: 13.5pt; color: #3730a3; margin: 20px 0 7px;
    padding-left: 10px; border-left: 4px solid #a5b4fc;
  }
  h3 { font-size: 11.5pt; color: #4338ca; margin: 15px 0 5px; }
  h1, h2, h3 { page-break-after: avoid; font-weight: 700; }
  p { margin: 7px 0; }
  strong { color: #111827; }
  a { color: #4f46e5; text-decoration: none; }
  ul, ol { margin: 7px 0; padding-left: 22px; }
  li { margin: 3px 0; }
  li::marker { color: #6366f1; }
  /* Citations */
  blockquote {
    border-left: 4px solid #c7d2fe; background: #f5f6ff; margin: 12px 0;
    padding: 6px 14px; color: #4b5563; border-radius: 0 6px 6px 0;
  }
  /* Code */
  code { background: #eef2ff; color: #3730a3; padding: 1px 5px; border-radius: 3px; font-size: 9pt; }
  pre { background: #1e1b4b; color: #e0e7ff; padding: 12px 14px; border-radius: 8px; font-size: 8.5pt; overflow-x: auto; }
  pre code { background: none; color: inherit; padding: 0; }
  /* Tableaux zébrés */
  table { border-collapse: collapse; width: 100%; margin: 14px 0; font-size: 9.5pt; page-break-inside: avoid; }
  th { background: #4f46e5; color: #fff; font-weight: 600; text-align: left; padding: 8px 11px; }
  td { padding: 7px 11px; border-bottom: 1px solid #e5e7eb; }
  tr:nth-child(even) td { background: #f8f8fc; }
  /* Séparateur (avant le bloc Sources) */
  hr { border: none; border-top: 1px solid #e5e7eb; margin: 22px 0 10px; }
  img { max-width: 100%; }
"""


def _html_document(titre: str, contenu_html: str) -> str:
    """Assemble le HTML complet stylé pour l'export PDF (rendu soigné du Markdown)."""
    from datetime import datetime
    date_str = datetime.now().strftime("%d/%m/%Y à %H:%M")
    titre_court = (titre or "Rapport").replace('"', "").strip()[:60]
    # `.replace()` et non l'opérateur `%` : le CSS contient des `100%` que `%` prendrait pour
    # des specificateurs de format (TypeError).
    css = _PDF_CSS.replace("__TITRE_COURT__", titre_court)
    return f"""<!DOCTYPE html>
<html lang="fr"><head><meta charset="UTF-8"><title>{titre}</title>
<style>{css}</style></head>
<body>
  <div class="doc-header">
    <h1>{titre}</h1>
    <div class="doc-meta">Matothèque · Rapport généré le {date_str}</div>
  </div>
  {contenu_html}
</body></html>"""


class ExportRequest(BaseModel):
    content: str = Field(..., min_length=1, description="Contenu Markdown à exporter")
    title: str = Field(default="Rapport DocFlow AI", description="Titre du document")


def _nom_export(title: str, extension: str) -> str:
    """Génère un nom de fichier propre pour l'export."""
    safe = "".join(c if c.isalnum() or c in "-_ " else "_" for c in title)
    safe = safe.strip().replace(" ", "_")[:50]
    horodatage = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{safe}_{horodatage}.{extension}"


@router.post("/export/pdf")
async def export_pdf(request: ExportRequest):
    """
    Convertit du Markdown en PDF et retourne le fichier.
    Utilise weasyprint via HTML comme intermédiaire.
    """
    try:
        from weasyprint import HTML
    except ImportError:
        raise HTTPException(status_code=500, detail="weasyprint non installé")

    nom_fichier = _nom_export(request.title, "pdf")

    # Convertir Markdown → HTML (extensions : tableaux, code, listes propres, retours à la ligne).
    contenu_html = markdown.markdown(
        request.content,
        extensions=["tables", "fenced_code", "nl2br", "sane_lists"],
    )
    html_complet = _html_document(request.title, contenu_html)

    try:
        # write_pdf() sans cible RETOURNE les octets → aucune écriture disque.
        pdf_bytes = HTML(string=html_complet).write_pdf()
    except Exception as e:
        log.error("Erreur génération PDF", erreur=str(e), type_err=type(e).__name__)
        raise HTTPException(status_code=500, detail=f"Erreur génération PDF : {e}")

    log.info("PDF généré", fichier=nom_fichier, octets=len(pdf_bytes))
    return _fichier_reponse(pdf_bytes, nom_fichier, "application/pdf")


@router.post("/export/docx")
async def export_docx(request: ExportRequest):
    """
    Convertit du Markdown en DOCX et retourne le fichier.
    Conversion basique : titres, paragraphes, listes.
    """
    nom_fichier = _nom_export(request.title, "docx")

    try:
        doc = DocxDocument()

        # Titre principal
        titre = doc.add_heading(request.title, level=0)
        titre.style.font.size = Pt(20)

        # Traitement ligne par ligne du Markdown
        lignes = request.content.split("\n")
        i = 0
        while i < len(lignes):
            ligne = lignes[i]

            if ligne.startswith("### "):
                doc.add_heading(ligne[4:].strip(), level=3)
            elif ligne.startswith("## "):
                doc.add_heading(ligne[3:].strip(), level=2)
            elif ligne.startswith("# "):
                doc.add_heading(ligne[2:].strip(), level=1)
            elif ligne.startswith("- ") or ligne.startswith("* "):
                # Liste à puces
                texte = ligne[2:].strip()
                p = doc.add_paragraph(texte, style="List Bullet")
            elif ligne.startswith("> "):
                # Citation
                p = doc.add_paragraph(ligne[2:].strip())
                p.style = doc.styles["Intense Quote"] if "Intense Quote" in doc.styles else p.style
            elif ligne.strip() == "":
                # Ligne vide : espace entre paragraphes
                pass
            elif ligne.startswith("---") or ligne.startswith("==="):
                # Séparateur — ignorer
                pass
            else:
                # Paragraphe normal
                if ligne.strip():
                    doc.add_paragraph(ligne.strip())

            i += 1

        # Sauvegarde EN MÉMOIRE (BytesIO) → aucune écriture disque, aucun droit requis.
        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

    except Exception as e:
        log.error("Erreur génération DOCX", erreur=str(e), type_err=type(e).__name__)
        raise HTTPException(status_code=500, detail=f"Erreur génération DOCX : {e}")

    log.info("DOCX généré", fichier=nom_fichier, octets=len(docx_bytes))
    return _fichier_reponse(
        docx_bytes, nom_fichier,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
