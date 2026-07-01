"""
Service Template Filler — Remplissage de templates DOCX
========================================================
Utilise docxtpl (Jinja2) pour remplir les champs {{ champ }}
d'un template DOCX avec les valeurs extraites par le LLM.

Flux :
  1. Charger le template + détecter ses champs
  2. Construire le contexte documentaire (textes extraits)
  3. Demander au LLM de retourner un JSON {champ: valeur}
  4. Remplir avec docxtpl → sauvegarder dans storage/exports/
"""

import json
import re
import uuid
from pathlib import Path

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from config import get_settings
from logger import get_logger

log = get_logger(__name__)
settings = get_settings()

# Nombre max de chars de contexte documentaire
MAX_CONTEXT_CHARS = 60_000

PROMPT_REMPLISSAGE = """\
Tu dois remplir un modèle de document professionnel.

Voici les documents source à analyser :
{contexte}

Instructions supplémentaires : {instructions}

Tu dois retourner un objet JSON contenant les valeurs pour chaque champ du template.
Champs à remplir : {champs}

Règles :
- Retourne UNIQUEMENT un objet JSON valide, sans explication ni markdown
- Si un champ ne peut pas être rempli, utilise une chaîne vide ""
- Sois précis et professionnel dans les valeurs
- Adapte le style au type de document

Réponds avec uniquement le JSON :"""


def _extraire_json(texte: str) -> dict:
    """Extrait le premier objet JSON trouvé dans la réponse du LLM."""
    # Tentative 1 : JSON direct
    try:
        return json.loads(texte.strip())
    except json.JSONDecodeError:
        pass

    # Tentative 2 : bloc ```json ... ```
    match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", texte)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass

    # Tentative 3 : premier { ... } du texte
    match = re.search(r"\{[\s\S]*\}", texte)
    if match:
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            pass

    log.warning("Impossible d'extraire JSON du LLM", texte_debut=texte[:200])
    return {}


def _construire_contexte(docs: list) -> str:
    """Construit le contexte documentaire en concaténant les textes extraits."""
    parts = []
    chars_restants = MAX_CONTEXT_CHARS

    for doc in docs:
        texte = doc.texte_extrait or ""
        if not texte.strip():
            continue

        entete = f"\n--- Document : {doc.nom} ---\n"
        espace_dispo = chars_restants - len(entete) - 100
        if espace_dispo <= 0:
            break

        if len(texte) > espace_dispo:
            texte = texte[:espace_dispo] + "\n[...tronqué...]"

        parts.append(entete + texte)
        chars_restants -= len(entete) + len(texte)

    return "\n".join(parts) if parts else "Aucun texte extrait disponible."


class TemplateFiller:
    """Remplit les templates DOCX avec des données extraites par IA."""

    def __init__(self, ollama_service):
        self.ollama = ollama_service
        self.exports_dir = Path(settings.storage_exports)
        self.exports_dir.mkdir(parents=True, exist_ok=True)

    def detect_fields(self, template_path: Path) -> list[dict]:
        """
        Détecte les champs Jinja2/docxtpl dans un template DOCX.

        Returns:
            Liste de {nom, type, description}
        """
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(template_path))

            # Collecter tout le texte (paragraphes + cellules de tableau)
            textes = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        textes.append(cell.text)

            texte_complet = "\n".join(textes)

            # Extraire les {{ champ }}
            champs = re.findall(r"\{\{\s*(\w+)\s*\}\}", texte_complet)

            # Dédupliquer en préservant l'ordre
            vus: set[str] = set()
            champs_uniques = []
            for c in champs:
                if c not in vus:
                    vus.add(c)
                    champs_uniques.append({"nom": c, "type": "texte", "description": None})

            return champs_uniques

        except Exception as e:
            log.warning("Impossible de détecter les champs DOCX", erreur=str(e))
            return []

    async def fill(
        self,
        template_id: str,
        document_ids: list[str],
        instructions: str | None,
        model: str | None,
        db: AsyncSession,
    ) -> Path:
        """
        Remplit un template avec les données extraites des documents.

        Returns:
            Chemin vers le fichier DOCX rempli
        """
        from models.document import Document
        from models.template import Template

        # 1. Charger le template depuis la DB
        result = await db.execute(
            select(Template).where(Template.id == uuid.UUID(template_id))
        )
        template = result.scalar_one_or_none()
        if not template:
            raise ValueError(f"Template '{template_id}' non trouvé")

        if template.type != "docx":
            raise ValueError("Seuls les templates DOCX sont supportés pour le remplissage")

        template_path = Path(template.chemin_fichier)
        if not template_path.exists():
            raise FileNotFoundError(f"Fichier template introuvable : {template_path}")

        # 2. Charger les documents
        doc_uuids = [uuid.UUID(d) for d in document_ids]
        result = await db.execute(
            select(Document).where(Document.id.in_(doc_uuids))
        )
        docs = result.scalars().all()

        if not docs:
            raise ValueError("Aucun document trouvé pour les IDs fournis")

        # 3. Détecter les champs (depuis la DB ou re-détection)
        champs = template.champs or self.detect_fields(template_path)
        noms_champs = [c["nom"] for c in champs]

        if not noms_champs:
            raise ValueError("Aucun champ {{ ... }} détecté dans le template")

        log.info(
            "Remplissage template",
            template=template.nom,
            nb_docs=len(docs),
            champs=noms_champs,
        )

        # 4. Construire le contexte
        contexte = _construire_contexte(docs)

        # 5. Demander au LLM les valeurs
        prompt = PROMPT_REMPLISSAGE.format(
            contexte=contexte,
            instructions=instructions or "Aucune instruction supplémentaire",
            champs=json.dumps(noms_champs, ensure_ascii=False),
        )

        from services import runtime_config
        model_utilise = model or runtime_config.model_for("rapport")
        reponse = await self.ollama.generate(prompt, model=model_utilise)
        valeurs = _extraire_json(reponse)

        # S'assurer que tous les champs ont une valeur (même vide)
        for champ in noms_champs:
            if champ not in valeurs:
                valeurs[champ] = ""

        log.info("Valeurs extraites", nb_champs_remplis=sum(1 for v in valeurs.values() if v))

        # 6. Remplir avec docxtpl
        try:
            from docxtpl import DocxTemplate
            tpl = DocxTemplate(str(template_path))
            tpl.render(valeurs)
        except Exception as e:
            log.error("Erreur docxtpl", erreur=str(e))
            raise RuntimeError(f"Erreur remplissage template : {e}") from e

        # 7. Sauvegarder
        nom_export = f"{template_path.stem}_rempli_{uuid.uuid4().hex[:8]}.docx"
        chemin_sortie = self.exports_dir / nom_export
        tpl.save(str(chemin_sortie))

        log.info("Template rempli", fichier=nom_export)
        return chemin_sortie
